import streamlit as st
import os
import sys
import tempfile
import pandas as pd
import docx

# Add src folder to sys.path
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), "src"))

from detectors import PIIDetectorPipeline
from replacement import SyntheticReplacer
from docx_processor import DocxProcessor
from evaluation import (
    extract_all_document_elements,
    build_original_pii_inventory,
    run_known_source_regression_check,
    run_detector_based_residual_audit
)

# -----------------------------------------------------------------------------
# 1. PAGE CONFIGURATION & CUSTOM CSS THEME
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="PII Redaction & Security Audit Engine",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Premium Dark Mode / Glassmorphism Design System
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

    /* Global Page Styling */
    .main {
        background-color: #0b0f19;
        color: #f1f5f9;
        font-family: 'Inter', system-ui, -apple-system, sans-serif;
    }

    /* ============================================================
       SIDEBAR — COMPLETE DARK THEME FIX
       ============================================================ */
    [data-testid="stSidebar"] {
        background-color: #0d1322 !important;
        border-right: 1px solid rgba(255,255,255,0.08) !important;
    }

    /* All text in sidebar → white */
    [data-testid="stSidebar"] * {
        color: #e2e8f0 !important;
    }

    /* Override white-background input boxes → dark */
    [data-testid="stSidebar"] input,
    [data-testid="stSidebar"] textarea {
        background-color: #1e293b !important;
        color: #f1f5f9 !important;
        border: 1px solid rgba(99,102,241,0.5) !important;
        border-radius: 8px !important;
    }

    /* Multiselect & Selectbox container boxes */
    [data-testid="stSidebar"] [data-baseweb="select"] > div,
    [data-testid="stSidebar"] [data-baseweb="input"] > div,
    [data-testid="stSidebar"] [data-baseweb="base-input"] {
        background-color: #1e293b !important;
        border: 1px solid rgba(99,102,241,0.5) !important;
        color: #f1f5f9 !important;
        border-radius: 8px !important;
    }

    /* Multiselect dropdown menu list */
    [data-baseweb="menu"],
    [data-baseweb="popover"] {
        background-color: #1e293b !important;
        border: 1px solid rgba(99,102,241,0.5) !important;
    }
    [data-baseweb="menu"] li,
    [data-baseweb="menu"] li * {
        color: #e2e8f0 !important;
        background-color: #1e293b !important;
    }
    [data-baseweb="menu"] li:hover,
    [data-baseweb="menu"] li:hover * {
        background-color: #4f46e5 !important;
        color: #ffffff !important;
    }

    /* Selected tags/chips in multiselect */
    [data-testid="stSidebar"] [data-baseweb="tag"] {
        background-color: #4f46e5 !important;
        border-radius: 6px !important;
    }
    [data-testid="stSidebar"] [data-baseweb="tag"] * {
        color: #ffffff !important;
    }

    /* Number input stepper buttons */
    [data-testid="stSidebar"] [data-testid="stNumberInput"] button,
    [data-testid="stSidebar"] button:not([kind="primary"]):not([kind="secondary"]) {
        background-color: #2d3748 !important;
        color: #e2e8f0 !important;
        border-color: rgba(99,102,241,0.4) !important;
    }

    /* File uploader drag-drop box */
    [data-testid="stSidebar"] [data-testid="stFileUploadDropzone"],
    [data-testid="stSidebar"] [data-testid="stFileUploadDropzone"] * {
        background-color: #1e293b !important;
        border: 1px dashed rgba(99,102,241,0.6) !important;
        color: #cbd5e1 !important;
        border-radius: 10px !important;
    }

    /* Dividers */
    [data-testid="stSidebar"] hr {
        border-color: rgba(255,255,255,0.12) !important;
    }
    /* ============================================================ */

    
    /* Header Banner */
    .header-banner {
        background: linear-gradient(135deg, #1e1b4b 0%, #0f172a 50%, #1e293b 100%);
        border: 1px solid rgba(99, 102, 241, 0.3);
        border-radius: 20px;
        padding: 28px 36px;
        margin-bottom: 28px;
        box-shadow: 0 20px 30px -10px rgba(79, 70, 229, 0.2), 0 10px 15px -5px rgba(0, 0, 0, 0.5);
    }
    
    .header-title {
        font-size: 2.5rem;
        font-weight: 800;
        letter-spacing: -0.02em;
        background: linear-gradient(90deg, #818cf8 0%, #c084fc 50%, #38bdf8 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin: 0 0 8px 0;
    }
    
    .header-subtitle {
        color: #cbd5e1;
        font-size: 1.1rem;
        font-weight: 400;
        margin: 0;
    }
    
    /* Glassmorphic Metric Cards */
    .metric-card {
        background: rgba(15, 23, 42, 0.75);
        backdrop-filter: blur(16px);
        -webkit-backdrop-filter: blur(16px);
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 16px;
        padding: 24px;
        text-align: center;
        box-shadow: 0 8px 20px rgba(0, 0, 0, 0.3);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    }
    .metric-card:hover {
        transform: translateY(-4px);
        border-color: rgba(129, 140, 248, 0.6);
        box-shadow: 0 12px 28px rgba(99, 102, 241, 0.25);
    }
    
    .metric-value {
        font-size: 2.5rem;
        font-weight: 800;
        background: linear-gradient(135deg, #38bdf8 0%, #818cf8 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 6px;
    }
    
    .metric-label {
        font-size: 0.85rem;
        text-transform: uppercase;
        letter-spacing: 0.1em;
        color: #94a3b8;
        font-weight: 700;
    }

    /* Status Pill Badges */
    .status-badge-pass {
        background: linear-gradient(135deg, rgba(34, 197, 94, 0.2) 0%, rgba(16, 185, 129, 0.2) 100%);
        color: #4ade80;
        border: 1px solid rgba(34, 197, 94, 0.5);
        padding: 6px 16px;
        border-radius: 30px;
        font-weight: 700;
        font-size: 0.9rem;
        display: inline-block;
        box-shadow: 0 4px 12px rgba(34, 197, 94, 0.15);
    }
    .status-badge-fail {
        background: linear-gradient(135deg, rgba(239, 68, 68, 0.2) 0%, rgba(225, 29, 72, 0.2) 100%);
        color: #f87171;
        border: 1px solid rgba(239, 68, 68, 0.5);
        padding: 6px 16px;
        border-radius: 30px;
        font-weight: 700;
        font-size: 0.9rem;
        display: inline-block;
        box-shadow: 0 4px 12px rgba(239, 68, 68, 0.15);
    }

    /* Tabs Customization */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
        background-color: #0f172a;
        padding: 8px;
        border-radius: 16px;
        border: 1px solid rgba(255, 255, 255, 0.08);
    }

    .stTabs [data-baseweb="tab"] {
        height: 48px;
        border-radius: 10px;
        color: #94a3b8;
        font-weight: 600;
        font-size: 0.95rem;
        transition: all 0.2s ease;
    }

    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #4f46e5 0%, #6366f1 100%) !important;
        color: #ffffff !important;
        box-shadow: 0 4px 14px rgba(79, 70, 229, 0.4);
    }

    /* Buttons */
    .stButton>button {
        border-radius: 12px;
        font-weight: 700;
        font-size: 1rem;
        height: 48px;
        transition: all 0.2s ease;
    }
    .stDataFrame {
        border-radius: 12px;
        overflow: hidden;
        border: 1px solid rgba(255, 255, 255, 0.08);
    }
</style>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 2. CACHED RESOURCE INITIALIZATION
# -----------------------------------------------------------------------------
@st.cache_resource
def load_detector_pipeline():
    """Cache detector pipeline loading (spaCy NLP model + Regex engines)."""
    return PIIDetectorPipeline()

pipeline = load_detector_pipeline()

@st.cache_data
def get_cached_original_inventory(doc_path: str):
    """Caches original document PII inventory extraction for high-performance execution."""
    return build_original_pii_inventory(doc_path, pipeline)

# -----------------------------------------------------------------------------
# 3. SIDEBAR CONTROLS & CONFIGURATION
# -----------------------------------------------------------------------------
with st.sidebar:
    st.image("https://img.icons8.com/isometric-headers/100/security-shield.png", width=64)
    st.title("Scaler AI PII Redaction Tool")
    st.caption("Scaler AI Labs — Advanced PII Redaction Engine")
    st.markdown("---")

    # Document Selection Mode
    st.subheader("📄 Document Source")
    doc_source = st.radio(
        "Choose Input Document:",
        ["Red Herring Prospectus (Default)", "Upload Custom DOCX"],
        index=0
    )

    input_path = None
    input_filename = "Red Herring Prospectus.docx"

    if doc_source == "Red Herring Prospectus (Default)":
        default_file = os.path.join("input", "Red Herring Prospectus.docx")
        if os.path.exists(default_file):
            input_path = default_file
        else:
            st.error("Default `Red Herring Prospectus.docx` not found in `input/` directory.")
    else:
        uploaded_file = st.file_uploader("Upload `.docx` File", type=["docx"])
        if uploaded_file is not None:
            # Save uploaded file to temp path
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            input_path = temp_path
            input_filename = uploaded_file.name
            st.success(f"Loaded: `{uploaded_file.name}`")

    st.markdown("---")

    # Category Selection Filter
    st.subheader("⚙️ Active PII Categories")
    all_categories = [
        "PERSON", "EMAIL", "PHONE_NUMBER", "ORGANIZATION", "ADDRESS",
        "SSN", "CREDIT_CARD", "DATE_OF_BIRTH", "IP_ADDRESS", "PAN", "CIN"
    ]
    selected_categories = st.multiselect(
        "Detect & Redact Categories:",
        all_categories,
        default=all_categories
    )

    st.markdown("---")
    
    # Faker Synthetic Generation Controls
    st.subheader("🎲 Replacement Options")
    faker_seed = st.number_input("Deterministic Seed:", min_value=1, max_value=99999, value=42, step=1)
    
    st.markdown("---")

    # Action Triggers
    btn_redact = st.button("🚀 Run PII Redaction & Audit", type="primary", use_container_width=True)
    btn_dry_run = st.button("🔍 Run Dry-Run Analysis", use_container_width=True)

# -----------------------------------------------------------------------------
# 4. MAIN DASHBOARD CONTENT
# -----------------------------------------------------------------------------

# Header Banner
st.markdown("""
<div class="header-banner">
    <div class="header-title">🛡️ Scaler AI PII Redaction Tool</div>
    <div class="header-subtitle">Enterprise-grade document sanitization, deterministic synthetic replacement, & multi-layered residual audit</div>
</div>
""", unsafe_allow_html=True)

# Session State Initialization
if "redaction_complete" not in st.session_state:
    st.session_state.redaction_complete = False

# Execute Redaction Pipeline
if btn_redact:
    if not input_path or not os.path.exists(input_path):
        st.error("Please provide a valid DOCX document before executing.")
    else:
        with st.status("⏳ Processing Document — Please Wait...", expanded=True) as status_box:
            st.write("🔍 **Step 1/3**: Initializing PII Redaction Engine & Faker synthetic generator...")
            replacer = SyntheticReplacer(seed=int(faker_seed))
            processor = DocxProcessor(pipeline, replacer)
            
            output_dir = "output"
            os.makedirs(output_dir, exist_ok=True)
            output_filename = f"Redacted_{os.path.splitext(input_filename)[0]}.docx"
            output_path = os.path.join(output_dir, output_filename)

            st.write("📝 **Step 2/3**: Scanning paragraphs, tables, headers, footers & replacing PII...")
            redact_stats = processor.redact_document(input_path, output_path)
            orig_inventory_set = processor.original_inventory_set

            st.write("🛡️ **Step 3/3**: Running 2-Layer Residual Audit & Known-Source Verification...")
            detector_audit = run_detector_based_residual_audit(output_path, pipeline, orig_inventory_set, replacer)
            regression_results = run_known_source_regression_check(output_path, pipeline=pipeline)
            regression_passed = all(r["passed"] for r in regression_results)

            status_box.update(label="✅ Processing & Residual Audit Complete!", state="complete", expanded=False)


        # Store in Session State
        st.session_state.redact_stats = redact_stats
        st.session_state.detector_audit = detector_audit
        st.session_state.regression_results = regression_results
        st.session_state.regression_passed = regression_passed
        st.session_state.orig_inventory_set = orig_inventory_set
        st.session_state.mapping_cache = replacer.mapping_cache
        st.session_state.output_path = output_path
        st.session_state.output_filename = output_filename
        st.session_state.input_path = input_path
        st.session_state.redaction_complete = True

        st.toast("PII Redaction & Residual Audit Successfully Completed!", icon="✅")

# Execute Dry Run Analysis
if btn_dry_run:
    if not input_path or not os.path.exists(input_path):
        st.error("Please provide a valid DOCX document before executing dry-run.")
    else:
        with st.spinner("Scanning document for PII entities (Dry-Run Mode)..."):
            processor = DocxProcessor(pipeline, SyntheticReplacer(seed=42))
            counts = processor.dry_run(input_path)
            st.session_state.dry_run_counts = counts
            st.session_state.show_dry_run = True

# Display Dry-Run Results Modal / Expander if triggered
if getattr(st.session_state, "show_dry_run", False):
    st.info("### 🔍 Dry-Run Entity Detection Summary")
    df_dry = pd.DataFrame(list(st.session_state.dry_run_counts.items()), columns=["PII Entity Category", "Detected Count"])
    st.dataframe(df_dry, use_container_width=True)

# -----------------------------------------------------------------------------
# 5. DASHBOARD METRICS & TABBED INTERFACE
# -----------------------------------------------------------------------------
if st.session_state.redaction_complete:
    stats = st.session_state.redact_stats
    audit = st.session_state.detector_audit
    reg_pass = st.session_state.regression_passed
    audit_pass = (audit["original_pii_leaks"] == 0 and audit["new_or_unmatched_pii_like"] == 0)

    # Top Metric Banner Cards
    m1, m2, m3, m4, m5 = st.columns(5)
    
    with m1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['total_redactions']}</div>
            <div class="metric-label">Total PII Redactions</div>
        </div>
        """, unsafe_allow_html=True)
        
    with m2:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['paragraphs_processed']}</div>
            <div class="metric-label">Paragraphs Processed</div>
        </div>
        """, unsafe_allow_html=True)
        
    with m3:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['tables_processed']}</div>
            <div class="metric-label">Tables Audited</div>
        </div>
        """, unsafe_allow_html=True)
        
    with m4:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{stats['cells_processed']}</div>
            <div class="metric-label">Table Cells Audited</div>
        </div>
        """, unsafe_allow_html=True)
        
    with m5:
        badge = '<span class="status-badge-pass">✅ PASS (0 LEAKS)</span>' if (audit_pass and reg_pass) else '<span class="status-badge-fail">❌ REVIEW</span>'
        st.markdown(f"""
        <div class="metric-card">
            <div style="margin-bottom: 8px;">{badge}</div>
            <div class="metric-label">Residual Audit Status</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # 5 Tab Navigation System
    t1, t2, t3, t4, t5 = st.tabs([
        "📊 Category & Benchmark Overview",
        "🔄 Synthetic Replacement Inspector",
        "📄 Document Preview & Snippets",
        "🛡️ Post-Redaction Residual Audit",
        "💾 Export Redacted DOCX"
    ])

    # -------------------------------------------------------------------------
    # TAB 1: CATEGORY & BENCHMARK OVERVIEW
    # -------------------------------------------------------------------------
    with t1:
        st.subheader("📈 Redactions Breakdown by PII Category")
        
        c1, c2 = st.columns([1.2, 1])
        
        with c1:
            by_type = stats.get("redactions_by_type", {})
            if by_type:
                df_cat = pd.DataFrame(list(by_type.items()), columns=["Category", "Redactions"]).sort_values("Redactions", ascending=False)
                st.bar_chart(df_cat, x="Category", y="Redactions", color="#3b82f6")
            else:
                st.info("No PII redactions recorded.")
                
        with c2:
            st.markdown("#### 🎯 Controlled Benchmark Evaluation Metrics")
            st.markdown("""
            Quantitative evaluation performed against the ground-truth benchmark suite ([`test_cases.json`](file:///d:/Downloads/Scaler/evaluation/test_cases.json)):
            """)
            
            b1, b2 = st.columns(2)
            with b1:
                st.metric("Precision", "83.87%", delta="High Quality")
                st.metric("F1-Score", "88.14%", delta="Balanced")
            with b2:
                st.metric("Recall", "92.86%", delta="Target >90% Met")
                st.metric("Accuracy", "82.05%", delta="Benchmark")

        st.markdown("---")
        st.subheader("📋 Ground-Truth Benchmark Per-Type Performance")
        
        benchmark_data = [
            {"Category": "PERSON", "TP": 8, "FP": 1, "FN": 1, "Precision": "88.9%", "Recall": "88.9%", "F1": "88.9%", "Status": "Validated"},
            {"Category": "EMAIL", "TP": 4, "FP": 0, "FN": 0, "Precision": "100.0%", "Recall": "100.0%", "F1": "100.0%", "Status": "Validated"},
            {"Category": "PHONE_NUMBER", "TP": 5, "FP": 2, "FN": 0, "Precision": "71.4%", "Recall": "100.0%", "F1": "83.3%", "Status": "Validated"},
            {"Category": "ORGANIZATION", "TP": 1, "FP": 1, "FN": 0, "Precision": "50.0%", "Recall": "100.0%", "F1": "66.7%", "Status": "Validated"},
            {"Category": "ADDRESS", "TP": 2, "FP": 1, "FN": 1, "Precision": "66.7%", "Recall": "66.7%", "F1": "66.7%", "Status": "Validated"},
            {"Category": "SSN", "TP": 1, "FP": 0, "FN": 0, "Precision": "100.0%", "Recall": "100.0%", "F1": "100.0%", "Status": "Validated"},
            {"Category": "CREDIT_CARD", "TP": 1, "FP": 0, "FN": 0, "Precision": "100.0%", "Recall": "100.0%", "F1": "100.0%", "Status": "Validated"},
            {"Category": "DATE_OF_BIRTH", "TP": 1, "FP": 0, "FN": 0, "Precision": "100.0%", "Recall": "100.0%", "F1": "100.0%", "Status": "Context Enforced"},
            {"Category": "IP_ADDRESS", "TP": 1, "FP": 0, "FN": 0, "Precision": "100.0%", "Recall": "100.0%", "F1": "100.0%", "Status": "Validated"},
            {"Category": "PAN", "TP": 1, "FP": 0, "FN": 0, "Precision": "100.0%", "Recall": "100.0%", "F1": "100.0%", "Status": "India Extension"},
            {"Category": "CIN", "TP": 1, "FP": 0, "FN": 0, "Precision": "100.0%", "Recall": "100.0%", "F1": "100.0%", "Status": "India Extension"},
        ]
        st.dataframe(pd.DataFrame(benchmark_data), use_container_width=True)

    # -------------------------------------------------------------------------
    # TAB 2: SYNTHETIC REPLACEMENT INSPECTOR
    # -------------------------------------------------------------------------
    with t2:
        st.subheader("🔄 Deterministic Replacement Mapping Inventory")
        st.caption("Inspect how sensitive original PII entities were mapped to realistic Faker replacements during this run.")
        
        cache = st.session_state.mapping_cache
        if cache:
            rows = []
            for (orig_text, label), synth_val in cache.items():
                rows.append({
                    "Category": label,
                    "Original Real PII": orig_text,
                    "Synthetic Replacement (Faker)": synth_val
                })
            df_cache = pd.DataFrame(rows)
            
            search_query = st.text_input("🔍 Search Mappings:", placeholder="Type name, email, company, phone...")
            if search_query:
                df_cache = df_cache[
                    df_cache["Original Real PII"].str.contains(search_query, case=False, na=False) |
                    df_cache["Synthetic Replacement (Faker)"].str.contains(search_query, case=False, na=False) |
                    df_cache["Category"].str.contains(search_query, case=False, na=False)
                ]
            
            st.dataframe(df_cache, use_container_width=True, height=450)
            st.caption(f"Total Unique Entity Mappings Cached: **{len(rows)}**")
        else:
            st.info("No replacement mappings recorded.")

    # -------------------------------------------------------------------------
    # TAB 3: DOCUMENT PREVIEW & SNIPPETS
    # -------------------------------------------------------------------------
    with t3:
        st.subheader("📄 Paragraph & Table Snippet Comparison")
        st.caption("Preview snippets extracted from the document traversal.")
        
        try:
            orig_doc = docx.Document(st.session_state.input_path)
            red_doc = docx.Document(st.session_state.output_path)
            
            p_orig = [p.text for p in orig_doc.paragraphs if p.text.strip()]
            p_red = [p.text for p in red_doc.paragraphs if p.text.strip()]
            
            max_preview = min(len(p_orig), len(p_red), 15)
            
            st.markdown(f"Displaying top **{max_preview}** non-empty body paragraphs:")
            
            for i in range(max_preview):
                with st.expander(f"Paragraph {i+1} Snippet Preview"):
                    c_orig, c_red = st.columns(2)
                    with c_orig:
                        st.markdown("**Original Text:**")
                        st.info(p_orig[i][:400] + ("..." if len(p_orig[i]) > 400 else ""))
                    with c_red:
                        st.markdown("**Redacted Text:**")
                        st.success(p_red[i][:400] + ("..." if len(p_red[i]) > 400 else ""))
        except Exception as e:
            st.error(f"Error loading snippet comparison: {e}")

    # -------------------------------------------------------------------------
    # TAB 4: POST-REDACTION RESIDUAL AUDIT
    # -------------------------------------------------------------------------
    with t4:
        st.subheader("🛡️ Layer 1: Known-Source PII Regression Audit")
        st.caption("Verifies that specifically targeted original sensitive PII entities are 100% absent in the output DOCX.")
        
        reg_df = pd.DataFrame(st.session_state.regression_results)
        reg_df["Status"] = reg_df["passed"].map(lambda x: "✅ ABSENT (PASS)" if x else "❌ LEAKED (FAIL)")
        st.dataframe(reg_df[["type", "target", "status", "Status"]], use_container_width=True)
        
        st.markdown("---")
        st.subheader("🛡️ Layer 2: Whole-Value Normalized Residual Audit Breakdown")
        st.caption("Re-scans the redacted DOCX with `PIIDetectorPipeline` across all body paragraphs, tables, headers, and footers.")
        
        audit_rows = [
            {"Classification": "ORIGINAL_PII_LEAKS", "Count": audit["original_pii_leaks"], "Description": "Original real PII values leaking into output (MUST BE 0 FOR PASS)"},
            {"Classification": "SYNTHETIC_REPLACEMENTS", "Count": audit["synthetic_replacements"], "Description": "Faker synthetic replacements detected & matched against mapping cache"},
            {"Classification": "NEW_OR_UNMATCHED_PII_LIKE", "Count": audit["new_or_unmatched_pii_like"], "Description": "Unmatched PII-shaped spans not in original inventory"},
            {"Classification": "TOTAL DETECTED SPANS", "Count": audit["total_detected_spans"], "Description": "Total PII-shaped detections across all document elements"}
        ]
        st.dataframe(pd.DataFrame(audit_rows), use_container_width=True)

    # -------------------------------------------------------------------------
    # TAB 5: EXPORT & DOWNLOAD CENTER
    # -------------------------------------------------------------------------
    with t5:
        st.subheader("💾 Download Sanitized Output & Reports")
        
        d1, d2 = st.columns(2)
        
        with d1:
            st.markdown("#### 📄 Redacted DOCX Document")
            st.caption("Download the final redacted document ready for distribution.")
            if os.path.exists(st.session_state.output_path):
                with open(st.session_state.output_path, "rb") as f:
                    st.download_button(
                        label="📥 Download Redacted DOCX",
                        data=f.read(),
                        file_name=st.session_state.output_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True
                    )
                    
        with d2:
            st.markdown("#### 📝 Markdown Evaluation Report")
            st.caption("Download the formal evaluation & post-redaction audit report.")
            eval_report_path = os.path.join("evaluation", "evaluation_report.md")
            if os.path.exists(eval_report_path):
                with open(eval_report_path, "r", encoding="utf-8") as f:
                    st.download_button(
                        label="📥 Download Audit Report (MD)",
                        data=f.read(),
                        file_name="evaluation_report.md",
                        mime="text/markdown",
                        use_container_width=True
                    )

else:
    # Initial State View
    st.info("👈 Please select your document in the sidebar and click **🚀 Run PII Redaction & Audit** to start!")
