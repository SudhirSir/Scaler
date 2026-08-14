"""
DOCX Processor Module - Safe Document Traversal & Run-Level Redaction
Author: Sudhir Singh
Description: Traverses python-docx paragraphs, table cells, headers, and footers.
Replaces detected PII spans at the run level while preserving font styles, bold, italic, and formatting.
"""

import docx
from typing import Dict, List, Set, Tuple
from detectors import PIIDetectorPipeline, PIISpan
from replacement import SyntheticReplacer

class DocxProcessor:
    """Safely processes and redacts Microsoft Word (.docx) documents with run-level style preservation."""

    def __init__(self, pipeline: PIIDetectorPipeline, replacer: SyntheticReplacer):
        self.pipeline = pipeline
        self.replacer = replacer

    def _redact_paragraph_runs(self, paragraph: docx.text.paragraph.Paragraph) -> List[PIISpan]:
        """
        Redacts PII spans in a docx Paragraph at the run level,
        preserving run-level formatting (bold, italic, font properties)
        and supporting cross-run entity spans using position-based character reconstruction.
        """
        if not paragraph.runs:
            if paragraph.text and paragraph.text.strip():
                spans = self.pipeline.detect(paragraph.text)
                if spans:
                    spans.sort(key=lambda s: s.start)
                    chunks = []
                    last_idx = 0
                    for s in spans:
                        if s.start < last_idx:
                            continue
                        chunks.append(paragraph.text[last_idx:s.start])
                        chunks.append(self.replacer.get_replacement(s.text, s.label))
                        last_idx = s.end
                    chunks.append(paragraph.text[last_idx:])
                    paragraph.text = "".join(chunks)
                    return spans
            return []

        # 1. Build concatenated full text and record run character offsets
        full_text_chunks = []
        run_offsets = []
        curr_offset = 0

        for r in paragraph.runs:
            t = r.text
            start = curr_offset
            end = curr_offset + len(t)
            run_offsets.append((start, end, r))
            full_text_chunks.append(t)
            curr_offset = end

        full_text = "".join(full_text_chunks)
        if not full_text or not full_text.strip():
            return []

        spans = self.pipeline.detect(full_text)
        if not spans:
            return []

        spans.sort(key=lambda s: s.start)

        # Filter non-overlapping spans
        resolved_spans = []
        last_idx = 0
        for span in spans:
            if span.start >= last_idx:
                resolved_spans.append(span)
                last_idx = span.end

        if not resolved_spans:
            return []

        # Pre-calculate synthetic replacements for resolved spans
        span_replacements = {}
        for s in resolved_spans:
            span_replacements[(s.start, s.end)] = self.replacer.get_replacement(s.text, s.label)

        def get_span_at(idx: int):
            for s in resolved_spans:
                if s.start <= idx < s.end:
                    return s
            return None

        # 2. Position-based character reconstruction per run
        for r_start, r_end, r_obj in run_offsets:
            chunks = []
            i = r_start
            while i < r_end:
                span = get_span_at(i)
                if span is None:
                    chunks.append(full_text[i])
                    i += 1
                else:
                    if i == span.start:
                        chunks.append(span_replacements[(span.start, span.end)])
                    i = span.end
            r_obj.text = "".join(chunks)

        return resolved_spans

    def redact_document(self, input_path: str, output_path: str) -> Dict:
        print(f"Reading document: {input_path}", flush=True)
        doc = docx.Document(input_path)

        stats = {
            "paragraphs_processed": 0,
            "tables_processed": 0,
            "cells_processed": 0,
            "total_redactions": 0,
            "redactions_by_type": {}
        }

        # 1. Process Body Paragraphs
        print(f"Processing {len(doc.paragraphs)} body paragraphs...", flush=True)
        for p in doc.paragraphs:
            stats["paragraphs_processed"] += 1
            spans = self._redact_paragraph_runs(p)
            if spans:
                stats["total_redactions"] += len(spans)
                for s in spans:
                    stats["redactions_by_type"][s.label] = stats["redactions_by_type"].get(s.label, 0) + 1

        # 2. Process Body Tables with Merged Cell Deduplication
        print(f"Processing {len(doc.tables)} tables...", flush=True)
        for table in doc.tables:
            stats["tables_processed"] += 1
            processed_cells = set()
            for row in table.rows:
                for cell in row.cells:
                    tc_elem = cell._tc
                    if tc_elem in processed_cells:
                        continue
                    processed_cells.add(tc_elem)
                    stats["cells_processed"] += 1

                    for p in cell.paragraphs:
                        spans = self._redact_paragraph_runs(p)
                        if spans:
                            stats["total_redactions"] += len(spans)
                            for s in spans:
                                stats["redactions_by_type"][s.label] = stats["redactions_by_type"].get(s.label, 0) + 1

        # 3. Process Headers & Footers (Paragraphs & Header/Footer Tables)
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    spans = self._redact_paragraph_runs(p)
                    if spans:
                        stats["total_redactions"] += len(spans)
                        for s in spans:
                            stats["redactions_by_type"][s.label] = stats["redactions_by_type"].get(s.label, 0) + 1
                header_cells = set()
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tc_elem = cell._tc
                            if tc_elem in header_cells:
                                continue
                            header_cells.add(tc_elem)
                            for p in cell.paragraphs:
                                spans = self._redact_paragraph_runs(p)
                                if spans:
                                    stats["total_redactions"] += len(spans)
                                    for s in spans:
                                        stats["redactions_by_type"][s.label] = stats["redactions_by_type"].get(s.label, 0) + 1

            if section.footer:
                for p in section.footer.paragraphs:
                    spans = self._redact_paragraph_runs(p)
                    if spans:
                        stats["total_redactions"] += len(spans)
                        for s in spans:
                            stats["redactions_by_type"][s.label] = stats["redactions_by_type"].get(s.label, 0) + 1
                footer_cells = set()
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tc_elem = cell._tc
                            if tc_elem in footer_cells:
                                continue
                            footer_cells.add(tc_elem)
                            for p in cell.paragraphs:
                                spans = self._redact_paragraph_runs(p)
                                if spans:
                                    stats["total_redactions"] += len(spans)
                                    for s in spans:
                                        stats["redactions_by_type"][s.label] = stats["redactions_by_type"].get(s.label, 0) + 1

        print(f"Saving redacted document to: {output_path}", flush=True)
        doc.save(output_path)
        return stats

    def dry_run(self, input_path: str) -> Dict[str, int]:
        print(f"Executing Dry-Run analysis on: {input_path}", flush=True)
        doc = docx.Document(input_path)
        counts: Dict[str, int] = {
            "PERSON": 0, "EMAIL": 0, "PHONE_NUMBER": 0, "ORGANIZATION": 0,
            "ADDRESS": 0, "SSN": 0, "CREDIT_CARD": 0, "DATE_OF_BIRTH": 0, "IP_ADDRESS": 0
        }

        def _count_spans(text: str):
            if text and text.strip():
                spans = self.pipeline.detect(text)
                for s in spans:
                    counts[s.label] = counts.get(s.label, 0) + 1

        for p in doc.paragraphs:
            _count_spans(p.text)

        for table in doc.tables:
            processed_cells = set()
            for row in table.rows:
                for cell in row.cells:
                    tc_elem = cell._tc
                    if tc_elem in processed_cells:
                        continue
                    processed_cells.add(tc_elem)
                    for p in cell.paragraphs:
                        _count_spans(p.text)

        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    _count_spans(p.text)
            if section.footer:
                for p in section.footer.paragraphs:
                    _count_spans(p.text)

        return counts
